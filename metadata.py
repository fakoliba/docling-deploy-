"""
Advanced Document Metadata Extraction for Document Q&A Application

This module provides functionality for:
1. Extracting detailed metadata from various document formats
2. Parsing document structure information
3. Displaying advanced metadata in the interface
"""

import os
import re
import json
import time
import pandas as pd
import streamlit as st
from typing import Dict, List, Any, Optional, Tuple
import logging
from PyPDF2 import PdfReader, PdfFileReader
from docx import Document
import datetime

# Try to import PyMuPDF, but provide graceful degradation if it's not available
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    st.warning("PyMuPDF not installed. Advanced PDF metadata extraction will be limited. Install with 'pip install PyMuPDF'")

# Try to import dateutil
try:
    import dateutil.parser
    DATEUTIL_AVAILABLE = True
except ImportError:
    DATEUTIL_AVAILABLE = False
    st.warning("python-dateutil not installed. Date parsing functionality will be limited. Install with 'pip install python-dateutil'")

import PIL
from PIL import Image
from PIL.ExifTags import TAGS

# Setup logging
logger = logging.getLogger(__name__)

class MetadataExtractor:
    def __init__(self):
        """Initialize the metadata extractor"""
        pass
    
    def extract_metadata(self, file, file_path: str = None) -> Dict[str, Any]:
        """
        Extract metadata from a file based on its type
        
        Args:
            file: File object or binary data
            file_path: Optional file path
            
        Returns:
            Dictionary of metadata
        """
        if file_path and file_path.lower().endswith('.pdf'):
            return self.extract_pdf_metadata(file)
        elif file_path and file_path.lower().endswith('.docx'):
            return self.extract_docx_metadata(file)
        elif file_path and file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif')):
            return self.extract_image_metadata(file)
        else:
            # Default metadata
            return {
                "filename": file_path if file_path else "Unknown file",
                "extraction_date": datetime.datetime.now().isoformat(),
                "file_extension": os.path.splitext(file_path)[1] if file_path else "",
            }
    
    def extract_pdf_metadata(self, file) -> Dict[str, Any]:
        """
        Extract detailed metadata from a PDF
        
        Args:
            file: PDF file object
            
        Returns:
            Dictionary of PDF metadata
        """
        try:
            # Basic metadata from PdfReader
            pdf = PdfReader(file)
            metadata = {}
            
            # Get document info if available
            if pdf.metadata:
                for key, value in pdf.metadata.items():
                    # Clean up keys by removing leading slashes
                    clean_key = key.strip('/') if isinstance(key, str) else key
                    metadata[clean_key] = str(value)
            
            # Add page count
            metadata["page_count"] = len(pdf.pages)
            
            # Use PyMuPDF for advanced metadata if possible
            if PYMUPDF_AVAILABLE:
                try:
                    # Seek to beginning of file if it's a file-like object
                    if hasattr(file, 'seek'):
                        file.seek(0)
                    
                    # Open with PyMuPDF for additional metadata
                    doc = fitz.open(stream=file.read() if hasattr(file, 'read') else file, filetype="pdf")
                    
                    # Extract more detailed information
                    metadata["format"] = "PDF " + doc.metadata.get("format", "")
                    metadata["encrypted"] = doc.is_encrypted
                    metadata["form_fields"] = bool(doc.is_form_pdf)
                    metadata["has_links"] = any(page.links() for page in doc)
                    
                    # Extract page sizes
                    page_sizes = []
                    page_fonts = set()
                    image_count = 0
                    
                    for page in doc:
                        # Page size
                        width, height = page.rect.width, page.rect.height
                        page_sizes.append(f"{width:.1f}x{height:.1f}")
                        
                        # Count images
                        image_list = page.get_images()
                        image_count += len(image_list)
                        
                        # Extract fonts
                        for font in page.get_fonts():
                            if font:
                                font_name = font[3] if len(font) > 3 else "Unknown"
                                page_fonts.add(font_name)
                    
                    metadata["page_sizes"] = list(set(page_sizes))
                    metadata["fonts"] = list(page_fonts)
                    metadata["image_count"] = image_count
                    
                    doc.close()
                except Exception as e:
                    logger.warning(f"PyMuPDF extraction error: {str(e)}")
            else:
                # Add basic info without PyMuPDF
                metadata["note"] = "Advanced PDF metadata extraction limited (PyMuPDF not installed)"
            
            # Add creation and modification dates in ISO format
            for date_field in ["creation_date", "creationDate", "mod_date", "modDate"]:
                if date_field in metadata:
                    try:
                        if DATEUTIL_AVAILABLE:
                            dt = dateutil.parser.parse(metadata[date_field])
                            metadata[date_field + "_iso"] = dt.isoformat()
                    except:
                        pass
            
            return metadata
        
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
            return {"error": str(e)}
    
    def extract_docx_metadata(self, file) -> Dict[str, Any]:
        """
        Extract detailed metadata from a DOCX file
        
        Args:
            file: DOCX file object
            
        Returns:
            Dictionary of DOCX metadata
        """
        try:
            # Seek to beginning of file if it's a file-like object
            if hasattr(file, 'seek'):
                file.seek(0)
                
            doc = Document(file)
            metadata = {}
            
            # Core properties
            core_props = doc.core_properties
            metadata["title"] = core_props.title
            metadata["author"] = core_props.author
            metadata["subject"] = core_props.subject
            metadata["keywords"] = core_props.keywords
            metadata["category"] = core_props.category
            metadata["comments"] = core_props.comments
            metadata["created"] = core_props.created.isoformat() if core_props.created else None
            metadata["modified"] = core_props.modified.isoformat() if core_props.modified else None
            metadata["last_modified_by"] = core_props.last_modified_by
            metadata["revision"] = core_props.revision
            
            # Document statistics
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            # Count sections
            metadata["section_count"] = len(doc.sections)
            
            # Count and tabulate headings by level
            heading_counts = {}
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    heading_level = para.style.name.replace('Heading ', '')
                    heading_counts[heading_level] = heading_counts.get(heading_level, 0) + 1
            
            metadata["heading_counts"] = heading_counts
            
            # Count images (approximate - counting shapes)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            metadata["approximate_image_count"] = image_count
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return {"error": str(e)}
    
    def extract_image_metadata(self, file) -> Dict[str, Any]:
        """
        Extract detailed metadata from an image file
        
        Args:
            file: Image file object
            
        Returns:
            Dictionary of image metadata
        """
        try:
            # Seek to beginning of file if it's a file-like object
            if hasattr(file, 'seek'):
                file.seek(0)
            
            # Open image with PIL
            img = Image.open(file)
            metadata = {}
            
            # Basic properties
            metadata["format"] = img.format
            metadata["mode"] = img.mode
            metadata["width"], metadata["height"] = img.size
            metadata["resolution"] = img.info.get("dpi", None)
            
            # Extract EXIF data if available
            exif_data = {}
            if hasattr(img, '_getexif') and img._getexif():
                exif = img._getexif()
                for tag_id, value in exif.items():
                    tag = TAGS.get(tag_id, tag_id)
                    exif_data[tag] = str(value)
                
                # Extract standard EXIF fields
                metadata["make"] = exif_data.get("Make", None)
                metadata["model"] = exif_data.get("Model", None)
                metadata["date_time"] = exif_data.get("DateTimeOriginal", exif_data.get("DateTime", None))
                metadata["exposure_time"] = exif_data.get("ExposureTime", None)
                metadata["f_number"] = exif_data.get("FNumber", None)
                metadata["iso"] = exif_data.get("ISOSpeedRatings", None)
                metadata["focal_length"] = exif_data.get("FocalLength", None)
                
                # Store all EXIF data
                metadata["exif"] = exif_data
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting image metadata: {str(e)}")
            return {"error": str(e)}


def extract_and_store_document_metadata(file, file_path: str = None) -> Dict[str, Any]:
    """
    Extract and store metadata for a document
    
    Args:
        file: Document file object
        file_path: Path to the file
        
    Returns:
        Extracted metadata dictionary
    """
    extractor = MetadataExtractor()
    
    # Extract metadata
    metadata = extractor.extract_metadata(file, file_path)
    
    # Add timestamp
    metadata["extraction_timestamp"] = time.time()
    metadata["extraction_datetime"] = datetime.datetime.now().isoformat()
    
    return metadata


def render_metadata_ui(metadata: Dict[str, Any]):
    """
    Render metadata in a user-friendly format in Streamlit
    
    Args:
        metadata: Document metadata dictionary
    """
    with st.expander("Document Metadata", expanded=False):
        # File information
        st.subheader("Document Information")
        
        # Basic file info
        st.write("**Basic Information**")
        basic_info = {}
        for key in ["filename", "file_type", "file_size", "page_count", "format"]:
            if key in metadata:
                basic_info[key] = metadata[key]
        
        st.json(basic_info)
        
        # Author information
        author_info = {}
        for key in ["author", "creator", "Author", "title", "Title", "subject", "Subject", "producer", "Producer"]:
            if key in metadata and metadata[key]:
                author_info[key] = metadata[key]
        
        if author_info:
            st.write("**Author Information**")
            st.json(author_info)
        
        # Dates
        date_info = {}
        for key in ["creation_date_iso", "mod_date_iso", "created", "modified", "date_time"]:
            if key in metadata and metadata[key]:
                date_info[key.replace("_iso", "")] = metadata[key]
        
        if date_info:
            st.write("**Date Information**")
            st.json(date_info)
        
        # Document structure (if available)
        structure_info = {}
        for key in ["page_count", "paragraph_count", "table_count", "section_count", 
                    "heading_counts", "approximate_image_count", "image_count",
                    "width", "height", "page_sizes"]:
            if key in metadata and metadata[key] is not None:
                structure_info[key] = metadata[key]
        
        if structure_info:
            st.write("**Document Structure**")
            st.json(structure_info)
        
        # Technical details
        technical_info = {}
        for key in ["fonts", "encrypted", "form_fields", "has_links", "resolution", "mode"]:
            if key in metadata and metadata[key] is not None:
                technical_info[key] = metadata[key]
        
        if technical_info:
            st.write("**Technical Details**")
            st.json(technical_info)
        
        # Camera information for images
        camera_info = {}
        for key in ["make", "model", "exposure_time", "f_number", "iso", "focal_length"]:
            if key in metadata and metadata[key] is not None:
                camera_info[key] = metadata[key]
        
        if camera_info:
            st.write("**Camera Information**")
            st.json(camera_info)
        
        # Full metadata
        with st.expander("Raw Metadata", expanded=False):
            # Remove exif dictionary to avoid cluttering the display
            display_metadata = {k: v for k, v in metadata.items() if k != "exif"}
            st.json(display_metadata)
